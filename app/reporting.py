from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.chart import LineChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


REPORT_DIR = Path("reports")


def build_reports(payload: dict[str, Any]) -> dict[str, str]:
    """
    Orquesta la generación de reportes analíticos basados en la simulación.
    Realiza un post-procesamiento de datos para la toma de decisiones.
    """
    REPORT_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"reporte_parqueadero_udes_{stamp}_{uuid4().hex[:6]}"
    
    docx_path = REPORT_DIR / f"{base}.docx"
    xlsx_path = REPORT_DIR / f"{base}.xlsx"
    html_path = REPORT_DIR / f"{base}.html"
    
    # Inyectar métricas analíticas calculadas antes de escribir los archivos
    payload["analytics_engine"] = _compute_deep_analytics(payload)
    
    write_docx(docx_path, payload)
    write_xlsx(xlsx_path, payload)
    write_html(html_path, payload, docx_path.name, xlsx_path.name)
    
    return {"docx": docx_path.name, "xlsx": xlsx_path.name, "html": html_path.name}


def _compute_deep_analytics(payload: dict[str, Any]) -> dict[str, Any]:
    """
    Motor interno de diagnóstico operacional.
    Genera métricas de comparación cruzada simulado-teórico y estimaciones de optimización.
    """
    c = payload["config"]
    t = payload["theoretical"]
    s = payload["simulated"]
    
    # 1. Delta de Estabilidad (Divergencia entre modelo teórico y simulación real)
    delta_rejection = abs(s["rejection_probability"] - t["pk"])
    convergence_status = "Resultados muy estables" if delta_rejection < 0.03 else "Variabilidad normal" if delta_rejection < 0.07 else "Desviación temporal por fluctuación"
    
    # 2. Análisis del Estado del Tráfico
    traffic_intensity = t["offered_traffic"] / c["capacity"] if c["capacity"] > 0 else math.inf
    if traffic_intensity >= 1.2:
        regime = "Sobrecarga Crítica (La demanda supera drásticamente el espacio disponible)"
    elif traffic_intensity >= 0.95:
        regime = "Congestión Alta Inestable (Sensible a ráfagas de llegada en horas pico)"
    elif traffic_intensity >= 0.70:
        regime = "Ocupación Eficiente Optimizada (Flujo balanceado con margen de seguridad)"
    else:
        regime = "Baja utilización del espacio físico"

    # 3. Estimación de puestos óptimos para mitigar el rechazo al < 5% (Fórmula Erlang-B inversa)
    current_k = c["capacity"]
    target_pk = 0.05
    suggested_k = current_k
    if s["rejection_probability"] > target_pk and t["offered_traffic"] < 1000:
        a_offered = t["offered_traffic"]
        for test_k in range(max(1, int(a_offered)), current_k + 150):
            b = 1.0
            for n in range(1, test_k + 1):
                b = (a_offered * b) / (n + a_offered * b)
            if b <= target_pk:
                suggested_k = test_k
                break
    
    slots_needed_delta = max(0, suggested_k - current_k)
    
    # 4. Proyección de usuarios no atendidos al mes (Asumiendo 26 días operativos)
    lost_users_per_day = (s["rejected"] / c["duration_hours"]) * 14.0 if c["duration_hours"] > 0 else 0
    lost_users_monthly = int(lost_users_per_day * 26)

    return {
        "delta_rejection": delta_rejection,
        "convergence_status": convergence_status,
        "traffic_intensity": traffic_intensity,
        "regime": regime,
        "suggested_k": suggested_k,
        "slots_needed_delta": slots_needed_delta,
        "lost_users_monthly": lost_users_monthly,
        "efficiency_index": (s["accepted"] / s["arrivals"]) * 100 if s["arrivals"] > 0 else 100.0
    }


def write_html(path: Path, payload: dict[str, Any], docx_name: str, xlsx_name: str) -> None:
    c = payload["config"]
    t = payload["theoretical"]
    s = payload["simulated"]
    ae = payload["analytics_engine"]
    
    if ae['slots_needed_delta'] > 0:
        optimization_msg = f"la infraestructura requiere agregar un estimado de <strong>{ae['slots_needed_delta']} puestos físicos</strong>, pasando de la capacidad actual de K={c['capacity']} a una capacidad recomendada de <strong>K_óptimo={ae['suggested_k']}</strong>."
    else:
        optimization_msg = f"la infraestructura actual cuenta con la capacidad óptima necesaria (K={c['capacity']} puestos) para mantener el margen de seguridad operativa bajo los objetivos establecidos."

    findings_html = "".join(f"<li><span class=\"bullet-topic\">Diagnóstico:</span> {item}</li>" for item in payload["analysis"]["findings"])
    recs_html = "".join(f"<li><span class=\"bullet-topic\">Acción Recomendada:</span> {item}</li>" for item in payload["analysis"]["recommendations"])
    
    log_rows = "".join(
        f"<tr><td>{row['id']}</td>"
        f"<td><code>{row['control_card']}</code></td>"
        f"<td>{row['entry_clock']}</td>"
        f"<td>{row['exit_clock'] or '<span class=\"active-badge\">Aún en parqueadero</span>'}</td>"
        f"<td><span class=\"badge-status status-{row['status']}\">{_assignment_label(row['status'])}</span></td></tr>"
        for row in payload.get("moto_log", [])[:250]
    )

    path.write_text(
        f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Análisis de Simulación UDES</title>
  <style>
    :root {{
      --primary: #115e59; --primary-dark: #0f4c46; --bg-main: #f8fafc;
      --text-dark: #0f172a; --text-muted: #475569; --border: #e2e8f0;
      --accent-green: #10b981; --accent-red: #ef4444; --accent-gray: #64748b;
    }}
    body {{ margin: 0; font-family: 'Segoe UI', system-ui, sans-serif; background: var(--bg-main); color: var(--text-dark); line-height: 1.5; }}
    main {{ max-width: 1100px; margin: 0 auto; padding: 40px 24px; }}
    
    header {{ display: flex; justify-content: space-between; align-items: center; border-bottom: 2px solid var(--border); padding-bottom: 24px; margin-bottom: 32px; gap: 20px; }}
    .header-title h1 {{ margin: 0; font-size: 28px; color: var(--primary); font-weight: 800; letter-spacing: -0.5px; }}
    .header-title p {{ margin: 4px 0 0; color: var(--text-muted); font-size: 14px; }}
    .seed-badge {{ background: #ccfbf1; color: var(--primary-dark); padding: 4px 10px; border-radius: 12px; font-family: monospace; font-weight: bold; font-size: 13px; display: inline-block; margin-top: 6px; }}
    
    nav {{ display: flex; gap: 12px; }}
    .btn-download {{ display: inline-flex; align-items: center; padding: 10px 18px; background: var(--primary); color: white; text-decoration: none; border-radius: 8px; font-weight: 600; font-size: 14px; transition: background 0.2s; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
    .btn-download:hover {{ background: var(--primary-dark); }}
    .btn-secondary {{ background: #e2e8f0; color: var(--text-dark); }}
    .btn-secondary:hover {{ background: #cbd5e1; }}
    
    .grid-metrics {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr)); gap: 16px; margin-bottom: 32px; }}
    .card-metric {{ background: white; border: 1px solid var(--border); border-radius: 12px; padding: 20px; box-shadow: 0 1px 3px rgba(0,0,0,0.02); border-top: 4px solid var(--primary); }}
    .card-metric.alert {{ border-top-color: var(--accent-red); }}
    .card-metric .label {{ font-size: 12px; font-weight: 600; color: var(--text-muted); text-transform: uppercase; }}
    .card-metric .val {{ font-size: 26px; font-weight: 700; color: var(--text-dark); margin-top: 4px; }}
    
    .grid-tables {{ display: grid; grid-template-columns: 1fr 1fr; gap: 24px; margin-bottom: 32px; }}
    @media (max-width: 850px) {{ .grid-tables {{ grid-template-columns: 1fr; }} }}
    
    section {{ background: white; border: 1px solid var(--border); border-radius: 12px; padding: 24px; margin-bottom: 24px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); }}
    section h2 {{ margin: 0 0 16px; font-size: 18px; color: var(--primary-dark); border-bottom: 1px solid var(--border); padding-bottom: 8px; font-weight: 700; }}
    
    table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
    th, td {{ text-align: left; padding: 12px; border-bottom: 1px solid var(--border); }}
    th {{ background: #f1f5f9; color: var(--text-dark); font-weight: 600; }}
    .table-side th {{ width: 55%; color: var(--text-muted); font-weight: 500; background: transparent; }}
    
    .ai-insights {{ background: #f0fdfa; border-left: 4px solid var(--primary); padding: 16px; border-radius: 0 8px 8px 0; font-size: 14.5px; margin-bottom: 20px; }}
    .ai-insights p {{ margin: 0 0 8px 0; }}
    
    ul, ol {{ margin: 0; padding-left: 20px; }}
    li {{ margin-bottom: 10px; color: var(--text-dark); }}
    .bullet-topic {{ font-weight: 600; color: var(--primary-dark); }}
    
    .excel-view {{ border-radius: 8px; overflow: hidden; border: 1px solid #334155; box-shadow: 0 10px 15px -3px rgba(0,0,0,0.3); }}
    .excel-title {{ background: #1e293b; padding: 12px 16px; font-size: 13px; font-weight: 600; color: #38bdf8; border-bottom: 1px solid #334155; display: flex; justify-content: space-between; }}
    .excel-table {{ font-family: 'Consolas', monospace; font-size: 12.5px; background: #0f172a; color: #e2e8f0; }}
    .excel-table th {{ background: #1e293b; color: #38bdf8; border: 1px solid #334155; }}
    .excel-table td {{ border: 1px solid #1e293b; background: #0f172a; padding: 8px 12px; }}
    .excel-table tr:nth-child(even) td {{ background: #020617; }}
    
    .badge-status {{ display: inline-block; padding: 2px 8px; border-radius: 6px; font-size: 11px; font-weight: bold; text-transform: uppercase; }}
    .status-dentro {{ background: #d1fae5; color: #065f46; }}
    .status-salio {{ background: #f1f5f9; color: #475569; }}
    .status-rechazada {{ background: #fee2e2; color: #991b1b; }}
    .active-badge {{ color: var(--accent-green); font-style: italic; font-weight: 500; }}
  </style>
</head>
<body>
  <main>
    <header>
      <div class="header-title">
        <h1>Informe y Análisis Operacional del Parqueadero</h1>
        <p>Sistema de Gestión de Espacio e Infraestructura UDES</p>
        <span class="seed-badge">Semilla de Simulación: {c['seed']}</span>
      </div>
      <nav>
        <a href="/reports/{docx_name}?download=1" class="btn-download">Descargar MS Word</a>
        <a href="/reports/{xlsx_name}?download=1" class="btn-download btn-secondary">Descargar MS Excel</a>
      </nav>
    </header>

    <div class="grid-metrics">
      <div class="card-metric">
        <div class="label">Eficiencia de Ingreso</div>
        <div class="val">{ae['efficiency_index']:.2f}%</div>
      </div>
      <div class="card-metric alert">
        <div class="label">Tasa de Motos Rechazadas</div>
        <div class="val">{_pct(s['rejection_probability'])}</div>
      </div>
      <div class="card-metric">
        <div class="label">Ocupación Promedio</div>
        <div class="val">{_fmt(s['average_occupancy'])} motos/h</div>
      </div>
      <div class="card-metric">
        <div class="label">No Atendidos al Mes (Est.)</div>
        <div class="val">{ae['lost_users_monthly']} usuarios</div>
      </div>
    </div>

    <section>
      <h2>Resumen Ejecutivo del Sistema</h2>
      <div class="ai-insights">
        <p><strong>Conclusión General:</strong> El sistema de estacionamiento de motocicletas opera bajo un estado de <strong>{ae['regime']}</strong>. Al evaluar las fórmulas matemáticas frente a los eventos simulados, se observa un comportamiento con <strong>{ae['convergence_status']}</strong> (Diferencia: <code>{ae['delta_rejection']:.5f}</code>).</p>
        <p>Para estabilizar el servicio y reducir la tasa de motos rechazadas al 5.00% objetivo, {optimization_msg}</p>
      </div>
      <p>{payload['analysis']['summary']}</p>
    </section>

    <div class="grid-tables">
      <section>
        <h2>Comparativa del Modelo Matemático Teórico</h2>
        <table class="table-side">
          <thead><tr><th>Indicador Base</th><th>Cálculo Teórico</th></tr></thead>
          <tbody>
            <tr><th>Tráfico Solicitado Total</th><td>{_fmt(t['offered_traffic'])} erlangs</td></tr>
            <tr><th>Probabilidad de Parqueadero Vacío</th><td>{_pct(t['p0'])}</td></tr>
            <tr><th>Probabilidad de Rechazo Físico</th><td>{_pct(t['pk'])}</td></tr>
            <tr><th>Tasa de Entrada Efectiva</th><td>{_fmt(t['lambda_effective'])} motos/h</td></tr>
            <tr><th>Porcentaje de Utilización Teórica</th><td>{_pct(t['rho'])}</td></tr>
          </tbody>
        </table>
      </section>

      <section>
        <h2>Resultados Reales de la Simulación</h2>
        <table class="table-side">
          <thead><tr><th>Indicador de Simulación</th><th>Valor Registrado</th></tr></thead>
          <tbody>
            <tr><th>Total de Motos que Llegaron</th><td>{s['arrivals']} flujos</td></tr>
            <tr><th>Motos que Lograron Ingresar</th><td>{s['accepted']} unidades</td></tr>
            <tr><th>Motos Rechazadas por Parqueadero Lleno</th><td>{s['rejected']} unidades</td></tr>
            <tr><th>Porcentaje del Tiempo Lleno</th><td>{_pct(s['full_time_ratio'])} del tiempo</td></tr>
            <tr><th>Tiempo de Permanencia Promedio</th><td>{_fmt(s['average_stay'])} horas</td></tr>
          </tbody>
        </table>
      </section>
    </div>

    <section>
      <h2>Interpretación de Patrones de Datos</h2>
      <ul>{findings_html}</ul>
    </section>

    <section>
      <h2>Recomendaciones Estratégicas</h2>
      <ol>{recs_html}</ol>
    </section>

    <section>
      <h2>Registro de Auditoría de Eventos (Primeros 250 movimientos)</h2>
      <div class="excel-view">
        <div class="excel-title">
          <span>Libro: control_ingreso_tarjetas | Capacidad Actual: {c['capacity']} puestos</span>
          <span>Vista: Primeros 250 registros</span>
        </div>
        <table class="excel-table">
          <thead>
            <tr><th>ID Moto</th><th>Tarjeta de Control</th><th>Hora Entrada</th><th>Hora Salida</th><th>Estado de Asignación</th></tr>
          </thead>
          <tbody>
            {log_rows}
          </tbody>
        </table>
      </div>
    </section>
  </main>
</body>
</html>""",
        encoding="utf-8",
    )


def write_docx(path: Path, payload: dict[str, Any]) -> None:
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(15, 23, 42)

    # Título Principal
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("INFORME DE CAPACIDAD Y ANÁLISIS DEL PARQUEADERO")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(17, 94, 89)

    # Subtítulo y Semilla
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Comparativa: Fórmulas Teóricas vs Simulación Dinámica por Computadora\nSemilla utilizada en la simulación: {payload['config']['seed']}")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().add_run("—" * 60).font.color.rgb = RGBColor(226, 232, 240)

    # Resumen Ejecutivo
    ae = payload["analytics_engine"]
    doc.add_heading("1. Resumen Ejecutivo del Sistema", level=1)
    p_dictamen = doc.add_paragraph()
    p_dictamen.paragraph_format.left_indent = Inches(0.25)

    if ae['slots_needed_delta'] > 0:
        recommendation_text = f"Se recomienda incrementar en {ae['slots_needed_delta']} los puestos de estacionamiento para reducir la tasa de rechazo al límite ideal del 5.00%."
    else:
        recommendation_text = "No se requieren puestos adicionales; la infraestructura actual es suficiente para absorber la demanda dentro de los márgenes de tolerancia."
    
    run_dictamen = p_dictamen.add_run(
        f"DIAGNÓSTICO AUTOMÁTICO: Evaluando el comportamiento de los datos, el parqueadero se encuentra actualmente en un estado de "
        f"'{ae['regime']}'. Los resultados simulados presentan una '{ae['convergence_status']}' comparados con el modelo matemático. "
        f"Se proyecta que este nivel de congestión cause la pérdida de atención de {ae['lost_users_monthly']} usuarios al mes. "
        f"{recommendation_text}"
    )
    run_dictamen.font.italic = True

    # Parámetros iniciales
    doc.add_heading("2. Configuración Inicial de la Simulación", level=1)
    config = payload["config"]
    _table(doc, [
        ("Capacidad Física Actual (K)", f"{config['capacity']} puestos"),
        ("Tasa de Llegada de Motos (λ)", f"{config['arrival_rate']} motocicletas/hora"),
        ("Tasa de Salida / Servicio (μ)", f"{config['service_rate']} motos/hora"),
        ("Duración de la Simulación", f"{config['duration_hours']} horas continuas"),
        ("Semilla Estocástica Utilizada", str(config['seed'])),
    ], "Parámetro Evaluado")

    # Tabla de comparación
    doc.add_heading("3. Cuadro Comparativo: Teoría Matemática vs Simulación Real", level=1)
    doc.add_paragraph("A continuación se evalúan las fórmulas teóricas de equilibrio frente a los resultados reales arrojados por el simulador:")
    
    t = payload["theoretical"]
    s = payload["simulated"]
    
    comparison_rows = [
        ("Intensidad de Tráfico Solicitado (A)", _fmt(t["offered_traffic"]) + " erlangs", _fmt(s["arrivals"]/config["duration_hours"]/config["service_rate"]) + " erlangs"),
        ("Porcentaje de Motos Rechazadas (Pk)", _pct(t["pk"]), _pct(s["rejection_probability"])),
        ("Porcentaje de Utilización del Espacio (ρ)", _pct(t["rho"]), _pct(s["utilization"])),
        ("Ocupación Promedio de Puestos (L)", _fmt(t["average_occupancy"]) + " puestos", _fmt(s["average_occupancy"]) + " puestos"),
        ("Porcentaje de Tiempo Totalmente Lleno", "No aplica en teoría", _pct(s["full_time_ratio"])),
        ("Tiempo de Estancia Promedio", f"{1/config['service_rate']:.4f} horas", f"{s['average_stay']:.4f} horas"),
    ]
    _three_col_table(doc, comparison_rows)

    # Diagnósticos
    doc.add_heading("4. Interpretación de Datos y Hallazgos", level=1)
    doc.add_paragraph(payload["analysis"]["summary"])
    for item in payload["analysis"]["findings"]:
        doc.add_paragraph(item, style="List Bullet")

    # Recomendaciones
    doc.add_heading("5. Plan de Acción y Recomendaciones", level=1)
    for item in payload["analysis"]["recommendations"]:
        doc.add_paragraph(item, style="List Number")

    # Registros
    doc.add_heading("6. Anexo: Historial de Movimientos (Primeras 120 motos)", level=1)
    doc.add_paragraph(f"Listado cronológico de ingreso y salida de vehículos bajo la semilla {config['seed']}:")
    
    log_rows = [
        (
            str(row["id"]),
            row["control_card"],
            row["entry_clock"],
            row["exit_clock"] or "Aún adentro",
            _assignment_label(row["status"]),
        )
        for row in payload.get("moto_log", [])[:120]
    ]
    _wide_table(doc, ["ID Moto", "N° Tarjeta Control", "Hora Entrada", "Hora Salida", "Estado Final"], log_rows)

    # Ajustar estilos visuales de títulos
    for heading in doc.paragraphs:
        if heading.style.name.startswith("Heading"):
            for run in heading.runs:
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(15, 76, 70)

    doc.save(path)


def write_xlsx(path: Path, payload: dict[str, Any]) -> None:
    wb = Workbook()
    
    # --- HOJA 1 ---
    ws = wb.active
    ws.title = "Resumen de Rendimiento"
    ws.views.sheetView[0].showGridLines = True

    ws.merge_cells("A1:C1")
    ws["A1"] = f"REPORTE OPERACIONAL DE CAPACIDAD (SEMILLA: {payload['config']['seed']})"
    ws["A1"].font = Font(size=11, bold=True, color="FFFFFF", name="Segoe UI")
    ws["A1"].fill = PatternFill("solid", fgColor="115E59")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 35

    ws["A3"] = "INDICADOR CLAVE"
    ws["B3"] = "VALOR TEÓRICO (FÓRMULAS)"
    ws["C3"] = "VALOR SIMULADO (REAL)"
    
    for col_letter in ['A', 'B', 'C']:
        cell = ws[f"{col_letter}3"]
        cell.font = Font(bold=True, color="FFFFFF", name="Segoe UI", size=10)
        cell.fill = PatternFill("solid", fgColor="0F172A")
        cell.alignment = Alignment(horizontal="center")
    
    c = payload["config"]
    t = payload["theoretical"]
    s = payload["simulated"]
    ae = payload["analytics_engine"]

    metricas_excel = [
        ("Capacidad de Puestos (K)", c["capacity"], c["capacity"], "0"),
        ("Tasa de Entrada de Motos (λ)", c["arrival_rate"], s["arrivals"]/c["duration_hours"], "0.00"),
        ("Tasa de Salida / Desocupación (μ)", c["service_rate"], 1/s["average_stay"] if s["average_stay"]>0 else 0, "0.0000"),
        ("Tráfico Total Solicitado (A)", t["offered_traffic"], (s["arrivals"]/c["duration_hours"])*s["average_stay"], "0.00"),
        ("Porcentaje de Motos Rechazadas", t["pk"], s["rejection_probability"], "0.00%"),
        ("Porcentaje de Utilización del Espacio", t["rho"], s["utilization"], "0.00%"),
        ("Ocupación Promedio de Puestos", t["average_occupancy"], s["average_occupancy"], "0.00"),
        ("Porcentaje de Tiempo Totalmente Lleno", "N/A", s["full_time_ratio"], "0.00%"),
        ("Porcentaje de Éxito de Ingreso", 1.0 - t["pk"], ae["efficiency_index"]/100.0, "0.00%"),
        ("Total de Motos Rechazadas (Excedente)", "N/A", s["rejected"], "0"),
    ]

    for idx, (lbl, val_t, val_s, fmt) in enumerate(metricas_excel, start=4):
        ws.cell(row=idx, column=1, value=lbl)
        cell_t = ws.cell(row=idx, column=2, value=val_t)
        if isinstance(val_t, (int, float)): cell_t.number_format = fmt
            
        cell_s = ws.cell(row=idx, column=3, value=val_s)
        if isinstance(val_s, (int, float)): cell_s.number_format = fmt
            
    _format_sheet(ws, num_rows=len(metricas_excel)+3)

    start_box = len(metricas_excel) + 6
    ws.merge_cells(f"A{start_box}:C{start_box}")
    ws[f"A{start_box}"] = "ANÁLISIS DE MEJORA Y OPTIMIZACIÓN"
    ws[f"A{start_box}"].font = Font(bold=True, size=11, color="115E59", name="Segoe UI")
    
    ws[f"A{start_box+1}"] = "Estado del Tráfico Detectado:"
    ws[f"B{start_box+1}"] = ae["regime"]
    ws[f"A{start_box+2}"] = "Ampliación Recomendada:"
    ws[f"B{start_box+2}"] = f"{ae['slots_needed_delta']} puestos adicionales sugeridos"
    ws[f"A{start_box+3}"] = "Capacidad Óptima Calculada:"
    ws[f"B{start_box+3}"] = f"K_óptimo = {ae['suggested_k']} puestos"
    
    for row_offset in range(1, 4):
        ws[f"A{start_box+row_offset}"].font = Font(bold=True, name="Segoe UI", size=10, color="475569")
        ws[f"B{start_box+row_offset}"].font = Font(italic=True, name="Segoe UI", size=10, color="0F172A")

    # --- HOJA 2 ---
    ts = wb.create_sheet("Historial de Ocupacion")
    ts.views.sheetView[0].showGridLines = True
    ts.append(["Hora de la Simulación", "Ocupación Real", "Motos Ingresadas Acum.", "Motos Rechazadas Acum."])
    
    for point in payload["timeline"]:
        ts.append([point["t"], point["occupancy"], point["accepted"], point["rejected"]])
    _format_sheet(ts, is_log=False)
    
    for row in ts.iter_rows(min_row=2, max_row=ts.max_row, min_col=1, max_col=4):
        row[0].number_format = "0.00"
        row[1].number_format = "0"
        row[2].number_format = "0"
        row[3].number_format = "0"

    chart = LineChart()
    chart.title = "Evolución de la Ocupación en el Tiempo"
    chart.style = 13
    chart.y_axis.title = "Motos dentro del Parqueadero"
    chart.x_axis.title = "Tiempo (Horas)"
    
    data = Reference(ts, min_col=2, min_row=1, max_row=ts.max_row)
    cats = Reference(ts, min_col=1, min_row=2, max_row=ts.max_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 13
    chart.width = 22
    
    if chart.series:
        chart.series[0].graphicalProperties.line.solidFill = "115E59"
        chart.series[0].graphicalProperties.line.width = 25000
        
    ts.add_chart(chart, "F2")

    # --- HOJA 3 ---
    log = wb.create_sheet("Libro de Movimientos")
    log.views.sheetView[0].showGridLines = True
    log.append(["ID Moto", "Código Tarjeta Control", "Hora de Entrada", "Hora de Salida", "Estado de Asignación"])
    
    for row in payload.get("moto_log", []):
        log.append([
            row["id"],
            row["control_card"],
            row["entry_clock"],
            row["exit_clock"] if row["exit_clock"] else "AÚN ADENTRO",
            f"● {_assignment_label(row['status'])}",
        ])
        status_cell = log.cell(row=log.max_row, column=5)
        status_cell.font = Font(color=_assignment_color(row["status"]), bold=True, name="Consolas", size=11)
        
    _format_log_sheet(log)
    wb.save(path)


def _table(doc: Document, rows: list[tuple[str, Any]], key_title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.autofit = False
    
    hdr = table.rows[0].cells
    hdr[0].text = key_title
    hdr[1].text = "Valor"
    
    for cell in hdr:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '115E59')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
        cells[0].width = Inches(3.2)
        cells[1].width = Inches(3.2)


def _three_col_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Indicador"
    hdr[1].text = "Cálculo Teórico"
    hdr[2].text = "Resultado Simulación"
    
    for cell in hdr:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '0F172A')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    for r1, r2, r3 in rows:
        cells = table.add_row().cells
        cells[0].text = str(r1)
        cells[1].text = str(r2)
        cells[2].text = str(r3)


def _wide_table(doc: Document, headers: list[str], rows: list[tuple[Any, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), '1E293B')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)
                run.font.bold = True
                
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].paragraphs[0].runs[0].font.size = Pt(9)


def _format_sheet(ws, num_rows: int = 0, is_log: bool = False) -> None:
    thin_border = Border(
        left=Side(style='thin', color='E2E8F0'),
        right=Side(style='thin', color='E2E8F0'),
        top=Side(style='thin', color='E2E8F0'),
        bottom=Side(style='thin', color='E2E8F0')
    )
    max_r = num_rows if num_rows > 0 else ws.max_row
    for row in ws.iter_rows(min_row=1, max_row=max_r, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = Alignment(vertical="center", horizontal="left")
            cell.font = Font(name="Segoe UI", size=10, color="0F172A")
            if cell.row in (1, 3): continue
            cell.border = thin_border
            if cell.row > 3 and cell.row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="F8FAFC")
    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 32


def _format_log_sheet(ws) -> None:
    header_fill = PatternFill("solid", fgColor="1E293B")
    even_fill = PatternFill("solid", fgColor="0F172A")
    odd_fill = PatternFill("solid", fgColor="1E1E24")
    header_font = Font(bold=True, color="38BDF8", name="Consolas", size=11)
    body_font = Font(color="E2E8F0", name="Consolas", size=10)
    
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="center", horizontal="center")
            if cell.row == 1:
                cell.fill = header_fill
                cell.font = header_font
            else:
                cell.fill = even_fill if cell.row % 2 == 0 else odd_fill
                if cell.column != 5: cell.font = body_font
                    
    widths = [16, 26, 22, 22, 32]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.freeze_panes = "A2"


def _fmt(value: float) -> str:
    return f"{value:.4f}"


def _pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def _assignment_label(status: str) -> str:
    labels = {
        "dentro": "Asignada con éxito",
        "salio": "Liberada (Salió)",
        "rechazada": "Rechazada (Lleno)",
    }
    return labels.get(status, status)


def _assignment_color(status: str) -> str:
    colors = {
        "dentro": "10B981",
        "salio": "94A3B8",
        "rechazada": "EF4444",
    }
    return colors.get(status, "E2E8F0")